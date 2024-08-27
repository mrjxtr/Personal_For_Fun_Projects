from docx import Document

# Create a new Document for the DOCX file
doc = Document()

# Adding the content to the DOCX file
doc.add_heading(
    "Script: Brian Niccol: Meet Starbucks' New CEO – His $300 Million Journey to Success",
    0,
)

sections = [
    (
        "Hook (0:00 - 0:20)",
        "Camera fades in on the Starbucks logo.\n\nNarrator:\n\"Imagine taking the helm of one of the most iconic brands in the world, leading it into a new era of innovation and success. That's exactly what's happening at Starbucks, where the man behind Chipotle's incredible turnaround, Brian Niccol, is stepping up as the new CEO. But who is Brian Niccol, and how did he amass a $300 million fortune? Stick around, because we're about to dive into the incredible journey of Starbucks' new leader.\"\n\nQuick montage of Starbucks locations and Brian Niccol in a boardroom setting.",
    ),
    (
        "Introduction (0:20 - 1:00)",
        'Camera shifts to the narrator in a studio.\n\nNarrator:\n"Welcome back to Smart Money Sagas, where we explore the stories of those who’ve turned their dreams into fortunes. I’m [Your Name], and today we’re uncovering the story of Brian Niccol—the man who turned Chipotle around and is now set to lead Starbucks into its next chapter. If you’re into stories of financial success and smart money moves, make sure to subscribe and hit that notification bell so you never miss out."\n\nLower third graphic: "Subscribe for more success stories!"',
    ),
    (
        "In-Depth Value Proposition (1:00 - 5:00)",
        'Cut to a split screen of Starbucks and Chipotle logos.\n\nNarrator:\n"Brian Niccol\'s rise to success is nothing short of extraordinary. Before being named Starbucks’ CEO, Niccol made waves at Chipotle Mexican Grill. When he took over as CEO in 2018, Chipotle was in dire straits—plagued by food safety scandals and plummeting sales. But Niccol brought with him a wealth of experience from Taco Bell, where he had already revolutionized the fast-food industry with innovative marketing and menu strategies."\n\nB-roll footage of Chipotle restaurants and digital ordering platforms.\n\nNarrator:\n"At Chipotle, Niccol didn’t just stop at damage control; he spearheaded a transformation that brought the company back from the brink. He focused on digital innovation, introducing online ordering and delivery options that catered to a modern, convenience-driven audience. Niccol also launched new menu items like lifestyle bowls, appealing to health-conscious consumers—a move that redefined Chipotle’s brand."\n\nOn-screen text: "Digital innovation and smart menu strategies = success."\n\nNarrator:\n"These smart moves not only revived Chipotle’s fortunes but also significantly increased its stock value, contributing to Niccol’s impressive $300 million net worth. His strategic investments in tech-forward companies, including a stake in Beyond Meat, also played a crucial role in his financial success."\n\nShow a graphic of a rising stock chart with Niccol’s face next to it.',
    ),
    (
        "First Engagement Trigger (5:00 - 5:30)",
        'Cut back to the narrator.\n\nNarrator:\n"Now, before we dive deeper into how Brian Niccol plans to revolutionize Starbucks, I want to hear from you! What do you think is the most important factor in turning a company around? Drop your thoughts in the comments below, and let’s get a conversation going."\n\nEncourage interaction with on-screen text: "Comment below with your thoughts!"',
    ),
    (
        "Extended Story or Case Study (5:30 - 7:00)",
        'Cut to a close-up of the Starbucks logo with upbeat music.\n\nNarrator:\n"Taking the reins at Starbucks, Niccol is poised to apply his winning strategies on an even bigger stage. His vision for Starbucks includes expanding the beverage menu beyond traditional coffee offerings and integrating more personalized customer experiences through data-driven insights. He plans to utilize cutting-edge technology to streamline service delivery, reducing wait times, and improving the overall customer experience."\n\nB-roll of Starbucks app, mobile ordering, and in-store digital touchpoints.\n\nNarrator:\n"But it’s not just about technology. Niccol’s leadership style is known for fostering a collaborative and inclusive work environment. At Chipotle, he empowered employees with growth opportunities, and he’s expected to bring the same culture of development to Starbucks. His approach will likely energize the workforce and enhance customer service—a win-win for both employees and customers."\n\nShow testimonials or quotes from Chipotle employees praising Niccol’s leadership.',
    ),
    (
        "Second Engagement Trigger (7:00 - 7:30)",
        'Camera back on the narrator with an engaging smile.\n\nNarrator:\n"So, what’s your favorite Starbucks drink? Let me know in the comments! And while you\'re at it, hit that like button if you\'re excited to see where Brian Niccol takes Starbucks next."\n\nOn-screen text: "Like if you’re excited for Starbucks’ future!"',
    ),
    (
        "Detailed Conclusion (7:30 - 8:30)",
        'Cut to a montage of Starbucks stores and customers enjoying their drinks.\n\nNarrator:\n"Brian Niccol\'s journey from Taco Bell to Chipotle, and now to Starbucks, is a masterclass in smart leadership and strategic thinking. With his wealth of experience and a proven track record, Niccol is set to lead Starbucks into a new era of innovation and growth. His focus on technology, personalized customer experiences, and a strong company culture will likely make Starbucks not just a coffeehouse but a brand synonymous with innovation."\n\nDisplay key points on-screen: "Innovation, Growth, Strong Culture."\n\nNarrator:\n"If you’re looking to make smart moves in your own life, whether in business or personal finance, take a page out of Brian Niccol’s book: Invest in innovation, stay ahead of trends, and never underestimate the power of a strong team."',
    ),
    (
        "Final Call to Action (8:30 - 9:00)",
        'Close-up shot of the narrator.\n\nNarrator:\n"Thanks for watching Smart Money Sagas! If you enjoyed this video and want to hear more success stories, make sure to subscribe and turn on notifications. And don’t forget to share this video with anyone who could use some inspiration on their journey to success."\n\nOn-screen text: "Subscribe & Share!"',
    ),
    (
        "Outro with End Screen (9:00 - 10:00)",
        'Show the channel’s logo with upbeat music playing.\n\nNarrator:\n"Check out these other videos on financial success stories—we’ve got plenty more inspiring tales just waiting to be uncovered. I’m [Your Name], and I’ll see you in the next one."\n\nEnd screen with links to other videos and a subscribe button.',
    ),
]

# Add each section with its content to the document
for heading, content in sections:
    doc.add_heading(heading, level=1)
    doc.add_paragraph(content)

# Save the DOCX file
docx_path = r"D:\Admin Files\Desktop\UPWORK\SMS\1-script.docx"
doc.save(docx_path)